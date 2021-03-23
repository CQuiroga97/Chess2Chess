from docx import Document
from docx.shared import Inches


import os
from pathlib import Path
estructuras = ["while", "switch", "if ", "if("]
tipoVariable = ["int ", "string ", "ofstream ", "ifstream ", "char ", "float ", "time_t"]
condiciones = ["==", "<", ">", "!=", "<=", ">=", "!"]
operadores = ["=", "+", "*", "/"]
std = ["cout", "cin", "break", "}", "<<", ">>", "return"]
libreriasComunes = ["std", "string", "iostream", "using namespace"]
estructura = ""
tipo = ""
variables = []
librerias = []


def esMetodo(linea):
    if linea.find("(") != -1 and linea.find(")") != -1:
        return True
    return False

#Definicion de linea conteniendo un comentario
def esComentario(linea):
    if linea.find("//") != -1 or linea.find("/*") != -1:
        return True
    return False

#Definicion de linea conteniendo un comentario
def esEstructura(linea):
    for e in estructuras:
        if linea.find(e) != -1:
            if e!="switch":
                for i in condiciones:
                    if(linea.find(i) != -1):
                        #print(e + " encontrado")
                        return True
            else:
                #print(e + " encontrado")
                return True

            
    return False
    
#Definicion de linea conteniendo un caso
def esCaso(linea):
    if linea.find(":") != -1 and linea.find("case") != -1:
        return True
    return False

#Definicion de linea conteniendo una declaracion
def esDef(linea):
    for i in tipoVariable:
        if linea.find(i) != -1:
            #print(i + " encontrado")
            if linea.find(","):
                for v in linea.split(i)[1].split(","):
                    if v.find("\"") == -1:
                        variables.append(v.split("=")[0].split(";")[0].split("(")[0].strip())
            else:
                variables.append(linea.split(i)[1].split(";")[0].split("=")[0])

            return True
    return False

#Definicion de linea conteniendo una libreria
def esLib(linea):
    if linea.find("include") != -1:
        librerias.append(linea.split("<")[1].split(">")[0])
        return True
    if linea.find("using") != -1:
        librerias.append(linea.split("using")[1])
        return True
    return False

#Definicion de linea conteniendo una libreria
def esLineaCorriente(linea):
    if linea.find(";") == -1:
        return False
    for i in operadores:
        if linea.find(i) != -1:
            return True
    for i in std:
        if linea.find(i) != -1:
            return True
    return False

def calificar(ruta, link):
    nota = 8
    #ruta = "C:\\Users\\quiro\\OneDrive\\Documentos\\Python\\Calificador de C++\\SUMATIVA\\Miguel Montealegre,Laura Daza,Santiago Aguirre\\santiago aguirre\\diseño sumativa.cpp"
    cantCasos = 0
    cantCorch = 0
    errores = []
    file = open(ruta, "r")
    lineas = file.readlines()
    numLineas = 0
    cantMain = 0
    for contenido in lineas:
        numLineas += 1
        linea = contenido.strip()
        if linea:
            if not(esComentario(linea)):

                if linea.find("{") != -1:
                    cantCorch += 1

                if linea.find("}") != -1:
                    cantCorch -= 1
                    #print("Corchete cerrado encontrada en la linea " + str(numLineas))
                
                elif(esMetodo(linea)):
                    if linea.find("main") != -1:
                            cantMain += 1
                    if(esEstructura(linea)):
                        
                        None
                        #print("Estructura encontrada en la linea " + str(numLineas))
                    else:
                        None
                        #print("Funcion encontrada en la linea " + str(numLineas))
                elif(esCaso(linea)):
                    cantCasos += 1
                    #print("Caso encontrado en la linea " + str(numLineas))
                elif(esLib(linea)):
                    None
                    #print("Libreria encontrada en la linea " + str(numLineas))
                elif(esDef(linea)):
                    None
                    #print("Definicion encontrada en la linea " + str(numLineas))
                elif(esLineaCorriente(linea)):
                    None
                    #print("Linea encontrada en la linea " + str(numLineas))
                elif linea.find("{") != -1 or linea.find("do") != -1:
                    None    
                else:
                    errores.append(numLineas)
            

    print ("Total de variables: " + str(len(variables)))
    print ("Total de librerias: " + str(len(librerias)))
    print ("Total de casos: " + str(cantCasos))
    print("Cantidad de errores: " + str(len(errores)))
    
    #Escribiendo en el archivo de texto
    document = Document()
    document.add_heading('Reporte de la sumativa', 0)
    document.add_paragraph('Dentro de su suamtiva demostraron tener un buen nivel de programacion, haciendo uso de diferentes estructuras y variables dentro del mismo')

    document.add_heading('Errores de codigo', level=1)
    er = 0
    if len(errores) == 1:
        nota -= 1
        er = 1
        document.add_paragraph("Dentro del código encontre solo este pequeño error, esta en la linea " +str(errores[0]) + ", en la proxima sesion tendran tiempo para solucionarlo")
    if(len(errores) > 1):
        
        er = 1
        document.add_paragraph("Dentro del codigo detecte una serie de errores, que aunque esten presentes, son faciles de solucionar, por favor solucionen los errores en las sigueintes lineas durante la proxima sesion: ")
        for i in errores:
            nota -= 1
            document.add_paragraph("Linea "+str(i))
    if cantCorch > 1:
        nota -= 1
        er = 1
        document.add_paragraph("Los corchetes estan mal cerrados, cuentenlos y diagnostiquen en cual se equivocaron")
    if cantMain > 1:
        nota -= 1
        er = 1
        document.add_paragraph("Solo puede haber 1 main() en el codigo, se encontraron " + str(cantMain))
    if er == 0:
        document.add_paragraph("No encontré errores en el codigo, veo que es un codigo muy limpio, los felicito")
    document.add_heading('Uso de variables', level=1)
    ex = ""
    if len(variables) < 10:
        nota -= 1
        ex = " corto "
    elif len(variables) >= 10 and len(variables) < 22:
        ex = " normal "
    elif len(variables) >= 22:
        ex = " exesivo "
    document.add_paragraph("Encontre un total de " + str(len(variables)) + " variables en el proyecto, lo cual es algo " + ex + " para una aplicación como esta, la proxima sesion deben sustentar el uso de las sigueitnes variables: ")
    for i in variables:
        document.add_paragraph(i)

    document.add_heading('Librerias', level=1)
    sim = []
    for i in librerias:
        flag = 0
        for n in libreriasComunes:
            if i.find(n) != -1:
                flag = 1
        if flag == 0:
            sim.append(i)
            
    if(len(sim)>0):

        document.add_paragraph("Encontre un total de " + str(len(librerias)) + " librerias en el código, me gusta saber que hacen uso de otras librerias, la proxima sesion deben explicar el uso de las siguientes: ")
        for i in sim:
            document.add_paragraph(" - " + i)
    else:
        nota -= 1
        document.add_paragraph("Encontre un total de " + str(len(librerias)) + " librerias en el código, noto tambien que no usan librerias nuevas o diferentes a las ya vistas, durante la proxima sesion investigaran nuevas librerias que pudieron haber usado en su proyecto")
    
    
    document.add_heading('Switch', level=1)
    if(cantCasos < 5):
        document.add_paragraph("Recuerden que la cantidad de casos eran 4 modulos y un ultimo caso para cerrar el programa, por favor agregarlo")
    if(cantCasos == 5):
        document.add_paragraph("Hicieron uso de los casos que se pedian, sin embargo, vamos a intentar agregar un nuevo caso con lo que ustedes quieran")
    if(cantCasos > 5):
        document.add_paragraph("Hicieron uso de los casos que se pedian, pero aun con esto, agregaron mas, esto demuestra que dominan el uso de los switch")
    nota -= 1
    if nota < 4:
        nota = 4
    document.add_heading('Nota final de ' + str(nota), level=1)
    if nota == 4:
        document.add_paragraph("Deben esforzarse mas por cumplir con los requisitos de las sumativas y trabajar mejor en equipo")
    if nota > 4 and nota < 6:
        document.add_paragraph("Es una sumativa buena, sin embargo, segun todo lo que demsotraron en la misma, en la proxima pueden dar mas de ustedes")
    if nota >= 6:
        document.add_paragraph("Muy buena sumativa, se nota el trabajo y esfuerzo que le dieron al proyecto")

    document.add_heading('Lo que aprendi de mi sumativa', level=1)
    document.save(str(link) + '\\Reporte de sumativa.docx')
    variables.clear()
    librerias.clear()
    
    

def revisar(ruta, carpeta):

    for i in os.scandir(ruta):
        if i.is_dir():
            revisar(Path(str(ruta) + "\\" + i.name), i.name)
        if i.is_file():
            if(i.name.split(".")[1] == "cpp"):
                print("---------------------------------------------------------------------------------------------------------")
                print("Calificando archivo "+i.name+" en la carpeta "+carpeta)
                print("---------------------------------------------------------------------------------------------------------")
                calificar(str(ruta)+"\\"+i.name, ruta)
revisar(Path("C:\\Users\\quiro\\OneDrive\\Documentos\\Python\\Calificador de C++"), "Raiz")